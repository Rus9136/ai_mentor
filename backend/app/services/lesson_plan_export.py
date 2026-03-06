"""
DOCX export for lesson plans (QMJ).
"""
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def export_to_docx(plan_data: dict, context_data: dict) -> BytesIO:
    """Convert lesson plan JSON to DOCX document."""
    doc = Document()

    # Landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    header = plan_data.get("header", {})

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("ҚЫСҚАМЕРЗІМДІ ЖОСПАР (ҚМЖ)")
    run.bold = True
    run.font.size = Pt(14)

    # Header table
    header_rows = [
        ("Бөлім (раздел)", header.get("section", "")),
        ("Тақырып (тема)", header.get("topic", "")),
        ("Оқу мақсаты", header.get("learning_objective", "")),
        ("Сабақтың мақсаты", header.get("lesson_objective", "")),
        ("Құндылық", f"{header.get('monthly_value', '')} — {header.get('value_decomposition', '')}"),
    ]

    # Add context info
    subject = context_data.get("subject", "")
    grade = context_data.get("grade_level", "")
    if subject or grade:
        header_rows.insert(0, ("Пән / Сынып", f"{subject}, {grade}-сынып"))

    ht = doc.add_table(rows=len(header_rows), cols=2)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(ht)
    for i, (label, value) in enumerate(header_rows):
        _set_cell(ht.cell(i, 0), label, bold=True, width=Cm(6))
        _set_cell(ht.cell(i, 1), value)

    doc.add_paragraph()

    # Stages
    stages = plan_data.get("stages", [])
    for stage in stages:
        # Stage header
        p = doc.add_paragraph()
        run = p.add_run(f"{stage.get('name', '')} ({stage.get('name_detail', '')})")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(f"  —  {stage.get('duration_min', '')} мин")
        run.font.size = Pt(11)

        # Method info
        method_rows = [
            ("Әдіс (метод)", stage.get("method_name", "")),
            ("Мақсаты", stage.get("method_purpose", "")),
            ("Тиімділігі", stage.get("method_effectiveness", "")),
        ]
        mt = doc.add_table(rows=3, cols=2)
        _style_table(mt)
        for i, (label, value) in enumerate(method_rows):
            _set_cell(mt.cell(i, 0), label, bold=True, width=Cm(5))
            _set_cell(mt.cell(i, 1), value)

        # Activities table
        act_t = doc.add_table(rows=2, cols=2)
        _style_table(act_t)
        _set_cell(act_t.cell(0, 0), "Педагогтің іс-әрекеті", bold=True)
        _set_cell(act_t.cell(0, 1), "Оқушының іс-әрекеті", bold=True)
        _set_cell(act_t.cell(1, 0), stage.get("teacher_activity", ""))
        _set_cell(act_t.cell(1, 1), stage.get("student_activity", ""))

        # Assessment
        if stage.get("assessment"):
            p = doc.add_paragraph()
            run = p.add_run("Бағалау: ")
            run.bold = True
            p.add_run(stage["assessment"])

        # Differentiation
        if stage.get("differentiation"):
            p = doc.add_paragraph()
            run = p.add_run("Саралау: ")
            run.bold = True
            p.add_run(stage["differentiation"])

        # Resources
        if stage.get("resources"):
            p = doc.add_paragraph()
            run = p.add_run("Ресурстар: ")
            run.bold = True
            p.add_run(stage["resources"])

        # Tasks with descriptors
        tasks = stage.get("tasks", [])
        if tasks:
            p = doc.add_paragraph()
            run = p.add_run("Тапсырмалар:")
            run.bold = True

            for task in tasks:
                p = doc.add_paragraph()
                run = p.add_run(f"Тапсырма {task.get('number', '')}")
                run.bold = True
                run = p.add_run(f" ({task.get('total_score', 0)} балл)")

                task_t = doc.add_table(rows=2, cols=2)
                _style_table(task_t)
                _set_cell(task_t.cell(0, 0), "Мұғалім", bold=True)
                _set_cell(task_t.cell(0, 1), "Оқушы", bold=True)
                _set_cell(task_t.cell(1, 0), task.get("teacher_activity", ""))
                _set_cell(task_t.cell(1, 1), task.get("student_activity", ""))

                descriptors = task.get("descriptors", [])
                if descriptors:
                    dt = doc.add_table(rows=len(descriptors) + 1, cols=2)
                    _style_table(dt)
                    _set_cell(dt.cell(0, 0), "Дескриптор", bold=True)
                    _set_cell(dt.cell(0, 1), "Балл", bold=True, width=Cm(2))
                    for di, d in enumerate(descriptors):
                        _set_cell(dt.cell(di + 1, 0), d.get("text", ""))
                        _set_cell(dt.cell(di + 1, 1), str(d.get("score", "")))

        doc.add_paragraph()

    # Total score
    p = doc.add_paragraph()
    run = p.add_run(f"Жалпы балл: {plan_data.get('total_score', 0)}")
    run.bold = True
    run.font.size = Pt(12)

    # Differentiation block
    diff = plan_data.get("differentiation")
    if diff:
        p = doc.add_paragraph()
        run = p.add_run("САРАЛАУ (дифференциация)")
        run.bold = True
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.add_run(diff.get("approach", ""))

        diff_t = doc.add_table(rows=3, cols=2)
        _style_table(diff_t)
        _set_cell(diff_t.cell(0, 0), "A деңгей", bold=True)
        _set_cell(diff_t.cell(0, 1), diff.get("for_level_a", ""))
        _set_cell(diff_t.cell(1, 0), "B деңгей", bold=True)
        _set_cell(diff_t.cell(1, 1), diff.get("for_level_b", ""))
        _set_cell(diff_t.cell(2, 0), "C деңгей", bold=True)
        _set_cell(diff_t.cell(2, 1), diff.get("for_level_c", ""))

    # Health & Safety
    hs = plan_data.get("health_safety", "")
    if hs:
        p = doc.add_paragraph()
        run = p.add_run("Денсаулық және қауіпсіздік: ")
        run.bold = True
        p.add_run(hs)

    # Reflection
    reflection = plan_data.get("reflection_template", [])
    if reflection:
        p = doc.add_paragraph()
        run = p.add_run("Рефлексия:")
        run.bold = True
        for item in reflection:
            doc.add_paragraph(item, style="List Bullet")

    # Write to BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _style_table(table):
    """Apply borders and basic styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
    # Set borders via XML
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )
    from lxml import etree
    tbl_pr.append(etree.fromstring(borders))


def _set_cell(cell, text: str, bold: bool = False, width=None):
    """Set cell text with optional formatting."""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = bold
    if width:
        cell.width = width
